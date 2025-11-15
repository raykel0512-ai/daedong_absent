# app.py

import streamlit as st
from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# =========================
# 1. 문서 생성 함수
# =========================

def create_absent_doc(
    school_name,
    grade,
    classname,
    number,
    student_name,
    start_date,
    end_date,
    days,
    reason_text,
    attach_med_cert,
    attach_rx,
    attach_parent_opinion,
    attach_etc,
    attach_etc_text,
    today,
    parent_name,
    parent_relation,
    homeroom_name,
    absent_type,         # 질병, 감염병, 경조사, 기타 등
    admit_type,          # 인정, 미인정
    confirm_method,      # 서류, 전화, 기타
    health_symptom,
    health_opinion,
    health_absent_date
):
    doc = Document()

    # 전체 기본 스타일(폰트) 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10)

    # ------------------------------------------------------------------
    # [1] 결석신고서
    # ------------------------------------------------------------------
    # 제목
    p = doc.add_paragraph()
    run = p.add_run("결 석 신 고 서")
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("※ 「결석신고서」는 결석한 날로부터 3일 이내에 제출하여 학교의 승인을 받아야 합니다.")
    doc.add_paragraph("   [  ]에는 해당되는 곳에 √표를 합니다. 「담임교사 확인서」는 결석신고서를 바탕으로 담임교사가 작성합니다.")

    # 학교명
    if school_name:
        doc.add_paragraph(f"\n학교명: {school_name}")

    # 학생 기본 정보 표
    table = doc.add_table(rows=2, cols=8)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "학년"
    hdr[1].text = grade or ""
    hdr[2].text = "반"
    hdr[3].text = classname or ""
    hdr[4].text = "번호"
    hdr[5].text = number or ""
    hdr[6].text = "성명"
    hdr[7].text = student_name or ""

    # 두 번째 줄 여유 칸 (추가 정보 필요 시 사용)
    row = table.rows[1].cells
    row[0].text = "비고"
    row[1].merge(row[7])
    row[1].text = ""

    # 결석 기간
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("결석 기간: ").bold = True
    p.add_run(
        f"{start_date.year}년 {start_date.month}월 {start_date.day}일"
        f"  ~  {end_date.year}년 {end_date.month}월 {end_date.day}일"
        f"   (공휴일 제외 {days}일간)"
    )

    doc.add_paragraph("※ 결석 기간 중 공휴일 또는 학교 휴무일은 결석 일수에 포함하지 않습니다.")

    # 결석 사유
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("결석 사유").bold = True
    doc.add_paragraph(reason_text or " ")

    # 붙임 (증빙 서류 체크)
    def box(checked: bool) -> str:
        # 체크박스 느낌
        return "■" if checked else "□"

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("붙임(증빙 서류)").bold = True

    doc.add_paragraph(
        f"{box(attach_med_cert)} 진단서 또는 진료 확인서 (3일 이상인 경우 필수 첨부)"
    )
    doc.add_paragraph(
        f"{box(attach_rx)} 병원 처방전 또는 약봉투"
    )
    doc.add_paragraph(
        f"{box(attach_parent_opinion)} 보건결석 학부모 의견서"
    )
    line_etc = f"{box(attach_etc)} 기타 증빙 서류"
    if attach_etc_text:
        line_etc += f" ({attach_etc_text})"
    doc.add_paragraph(line_etc)

    doc.add_paragraph("※ 규정된 증빙서류를 첨부하지 않으면 ‘미인정(무단)’ 결석 처리될 수 있습니다.")

    # 결론 문구 + 서명
    doc.add_paragraph("")
    doc.add_paragraph("위와 같이 결석하고자/하였기에 보호자 연서로 신고합니다.")

    doc.add_paragraph(
        f"\n{today.year}년  {today.month}월  {today.day}일"
    )

    doc.add_paragraph(
        f"\n학생 성명: {student_name}              (서명 또는 인)"
    )
    doc.add_paragraph(
        f"보호자 성명: {parent_name}              (서명 또는 인)"
    )
    if parent_relation:
        doc.add_paragraph(f"(학생과의 관계: {parent_relation})")

    # 페이지 나누기
    doc.add_page_break()

    # ------------------------------------------------------------------
    # [2] 담임교사 확인서
    # ------------------------------------------------------------------
    p = doc.add_paragraph()
    run = p.add_run("담 임 교 사 확 인 서")
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "※ 담임교사는 결석신고서 및 첨부 서류를 확인하여 아래 사항을 기재합니다."
    )

    # 기본 정보
    table2 = doc.add_table(rows=2, cols=6)
    table2.style = "Table Grid"

    r0 = table2.rows[0].cells
    r0[0].text = "학년"
    r0[1].text = grade or ""
    r0[2].text = "반"
    r0[3].text = classname or ""
    r0[4].text = "성명"
    r0[5].text = student_name or ""

    r1 = table2.rows[1].cells
    r1[0].text = "결석 기간"
    r1[1].merge(r1[5])
    r1[1].text = (
        f"{start_date.year}년 {start_date.month}월 {start_date.day}일"
        f" ~ {end_date.year}년 {end_date.month}월 {end_date.day}일"
        f" (총 {days}일)"
    )

    doc.add_paragraph("")

    # 인정 / 미인정 / 사유 분류
    doc.add_paragraph("① 결석 종류 및 인정 여부")

    table3 = doc.add_table(rows=3, cols=4)
    table3.style = "Table Grid"

    t3r0 = table3.rows[0].cells
    t3r0[0].text = "결석 구분"
    t3r0[1].text = "질병 / 감염병 / 경조사 / 기타"
    t3r0[2].text = "출석 인정 여부"
    t3r0[3].text = "비고"

    t3r1 = table3.rows[1].cells
    t3r1[0].text = "구분"
    t3r1[1].text = absent_type or ""
    t3r1[2].text = admit_type or ""
    t3r1[3].text = ""

    t3r2 = table3.rows[2].cells
    t3r2[0].text = "확인 방법"
    t3r2[1].text = confirm_method or ""
    t3r2[2].merge(t3r2[3])
    t3r2[2].text = "예: 서류 확인, 전화 통화, 기타 등"

    doc.add_paragraph("")

    doc.add_paragraph(
        "※ 출석인정 결석 해당 여부는 학업성적관리규정 및 관련 법령에 따릅니다."
    )

    doc.add_paragraph(
        f"\n담임교사: {homeroom_name or ''}                     (서명 또는 인)"
    )

    doc.add_page_break()

    # ------------------------------------------------------------------
    # [3] 보건결석 보호자 의견서
    # ------------------------------------------------------------------
    p = doc.add_paragraph()
    run = p.add_run("보건결석 보호자 의견서")
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("※ 보건결석(질병 등)인 경우에만 작성합니다.")

    # 기본 정보
    table4 = doc.add_table(rows=4, cols=4)
    table4.style = "Table Grid"

    r0 = table4.rows[0].cells
    r0[0].text = "학년"
    r0[1].text = grade or ""
    r0[2].text = "반"
    r0[3].text = classname or ""

    r1 = table4.rows[1].cells
    r1[0].text = "번호"
    r1[1].text = number or ""
    r1[2].text = "학생 성명"
    r1[3].text = student_name or ""

    r2 = table4.rows[2].cells
    r2[0].text = "결석(또는 조퇴 등) 일자"
    r2[1].merge(r2[3])
    if health_absent_date:
        r2[1].text = f"{health_absent_date.year}년 {health_absent_date.month}월 {health_absent_date.day}일"
    else:
        r2[1].text = ""

    r3 = table4.rows[3].cells
    r3[0].text = "증상"
    r3[1].merge(r3[3])
    r3[1].text = health_symptom or ""

    doc.add_paragraph("")

    doc.add_paragraph("보호자 의견 (자필 작성 내용 입력란):")
    doc.add_paragraph(health_opinion or " ")

    doc.add_paragraph(
        "\n위와 같이 보건결석 사유가 있었음을 확인합니다."
    )

    doc.add_paragraph(
        f"\n{today.year}년  {today.month}월  {today.day}일"
    )

    doc.add_paragraph(
        f"\n보호자 성명: {parent_name}            (서명 또는 인)"
    )
    if parent_relation:
        doc.add_paragraph(f"(학생과의 관계: {parent_relation})")

    # 버퍼로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =========================
# 2. Streamlit UI
# =========================

st.title("결석 신고서 · 담임교사 확인서 · 보호자 의견서 자동 생성기")

st.markdown("학교 양식에 맞춘 **결석 관련 서류**를 한 번에 docx로 생성합니다.")

st.subheader("기본 정보")
school_name = st.text_input("학교명", value="")
col1, col2, col3, col4 = st.columns(4)
with col1:
    grade = st.text_input("학년", value="")
with col2:
    classname = st.text_input("반", value="")
with col3:
    number = st.text_input("번호", value="")
with col4:
    student_name = st.text_input("학생 성명", value="")

st.subheader("결석 정보")
start_date = st.date_input("결석 시작일", value=date.today())
end_date = st.date_input("결석 종료일", value=date.today())
days = st.number_input("결석 일수(공휴일 제외)", min_value=1, value=1)

reason_text = st.text_area("결석 사유(구체적으로 기재)", height=120)

st.subheader("증빙 서류 (붙임)")
attach_med_cert = st.checkbox("진단서 또는 진료 확인서 (3일 이상 시 필수)", value=False)
attach_rx = st.checkbox("병원 처방전 또는 약봉투", value=False)
attach_parent_opinion = st.checkbox("보건결석 학부모 의견서", value=False)
attach_etc = st.checkbox("기타 증빙 서류", value=False)
attach_etc_text = st.text_input("기타 증빙 서류 내용(선택)", value="")

st.subheader("보호자 정보")
col1, col2 = st.columns(2)
with col1:
    parent_name = st.text_input("보호자 성명", value="")
with col2:
    parent_relation = st.text_input("학생과의 관계 (예: 모, 부, 조부모 등)", value="")

today = st.date_input("작성일", value=date.today())

st.subheader("담임교사 확인서(담임 작성용 정보)")
homeroom_name = st.text_input("담임교사 성명(양식에 표시용)", value="")

absent_type = st.selectbox(
    "결석 구분 (담임교사 확인서용)",
    ["", "질병", "감염병", "경조사", "기타"],
    index=0,
)
admit_type = st.selectbox(
    "출석 인정 여부",
    ["", "출석 인정", "미인정"],
    index=0,
)
confirm_method = st.selectbox(
    "확인 방법",
    ["", "증빙 서류 확인", "전화 통화", "기타"],
    index=0,
)

st.subheader("보건결석 보호자 의견서 (보건결석인 경우만)")
health_absent_date = st.date_input(
    "보건결석(또는 조퇴·지각 등) 일자", value=date.today()
)
health_symptom = st.text_input("증상(예: 발열, 기침, 복통 등)", value="")
health_opinion = st.text_area("보호자 의견(자필로 쓸 내용 입력란)", height=120)

st.markdown("---")

if st.button("서류 일괄 생성 (docx 다운로드)"):
    buffer = create_absent_doc(
        school_name=school_name,
        grade=grade,
        classname=classname,
        number=number,
        student_name=student_name,
        start_date=start_date,
        end_date=end_date,
        days=days,
        reason_text=reason_text,
        attach_med_cert=attach_med_cert,
        attach_rx=attach_rx,
        attach_parent_opinion=attach_parent_opinion,
        attach_etc=attach_etc,
        attach_etc_text=attach_etc_text,
        today=today,
        parent_name=parent_name,
        parent_relation=parent_relation,
        homeroom_name=homeroom_name,
        absent_type=absent_type,
        admit_type=admit_type,
        confirm_method=confirm_method,
        health_symptom=health_symptom,
        health_opinion=health_opinion,
        health_absent_date=health_absent_date,
    )

    file_name = f"결석서류_{student_name or '학생'}.docx"
    st.download_button(
        label="결석 신고서 · 확인서 docx 다운로드",
        data=buffer,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
